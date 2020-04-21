from docx import Document
from docx2pdf import convert
import os

document = Document("template.docx")

template = "template.docx"


def getInput():
    bs = input("Business name: ")
    position = input("Position: ")
    return bs, position

def op(bs,position):
    for sections in document.paragraphs:
        if "%COMPANY%" in sections.text:
            # replace company with the company name
            sections.text = sections.text.replace("%COMPANY%", bs)
        if "%POSITION%" in sections.text:
            #replace position with the position name
            sections.text = sections.text.replace("%POSITION%", position)
        document.paragraphs.append(sections.text)

    fileName = "generated//"+bs +" - "+position +".docx"
    document.save(fileName)
    return fileName

def con(fileName):
    FN = fileName.replace(".docx",".pdf")
    convert(fileName, FN)
    os.remove(fileName)



bs,pos = getInput()
fn = op(bs,pos)
con(fn)
